import docx
import os

def info_update(doc,old_info, new_info):
    '''此函数用于批量替换需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''
    #读取段落中的所有run，找到需替换的信息进行替换
    for para in doc.paragraphs: 
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info) #替换信息
    #读取表格中的所有单元格，找到需替换的信息进行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info) #替换信息

# path="F:\PYTHON\批量文档编辑\报告"  # 文件夹路径
path = os.getcwd()
print(path)

files=[]

for file in os.listdir(path):
    if file.endswith(".docx"): #排除文件夹内的其它干扰文件，只获取word文件
        files.append(path+file)
        print(file)
        
# for file in files:
#     doc = docx.Document(file)
#     print(file)
#     info_update(doc,"系统名称","系统名字")
#     #doc.save("data/替换结果/{}".format(file.split("/")[-1]))
#     print("{}替换完成".format(file))
